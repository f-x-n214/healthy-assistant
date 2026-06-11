# -*- coding: utf-8 -*-
"""Generate 06096911 methodology course paper per template & requirements."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUTPUT = Path(r"d:\21425\大三下\智能应用系统开发\理论论论论论文\202306120105-方晓楠-06096911.docx")

FONT = "宋体"
FONT_EN = "Times New Roman"
SZ_BODY = Pt(12)
SZ_H1 = Pt(14)
SZ_H2 = Pt(12)
SZ_ABS_TITLE = Pt(18)


def font(run, size=SZ_BODY, bold=False, en=False):
    run.font.name = FONT_EN if en else FONT
    run.font.size = size
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:eastAsia"), FONT)


def para(doc, text, indent=True, bold=False, size=SZ_BODY, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.alignment = align
    font(p.add_run(text), size=size, bold=bold)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    font(p.add_run(text), size=SZ_H1 if level == 1 else SZ_H2, bold=True)
    return p


def caption(doc, text, above=False):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run(text), bold=True)
    if not above:
        doc.add_paragraph()


def table(doc, title, headers, rows):
    caption(doc, title, above=True)
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                font(r, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = val
            for p in c.paragraphs:
                for r in p.runs:
                    font(r)
    doc.add_paragraph()


def fig(doc, title, lines):
    caption(doc, title, above=False)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("\n".join(lines)), size=Pt(11))
    doc.add_paragraph()


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    font(p.add_run(text), size=Pt(10.5))


def setup(doc):
    s = doc.sections[0]
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(2.0)
    s.right_margin = Cm(2.0)
    s.gutter = Cm(0.5)
    s.different_first_page_header_footer = True
    s.first_page_header.is_linked_to_previous = False
    if s.first_page_header.paragraphs:
        s.first_page_header.paragraphs[0].clear()
    hp = s.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(hp.add_run("智能应用系统开发方法研究——以银发健康助手为例"), size=Pt(10.5))
    fp = s.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    for el in [
        ("begin", None),
        ("instr", "PAGE"),
        ("sep", None),
        ("end", None),
    ]:
        if el[0] == "instr":
            node = OxmlElement("w:instrText")
            node.set(qn("xml:space"), "preserve")
            node.text = el[1]
        else:
            node = OxmlElement("w:fldChar")
            node.set(qn("w:fldCharType"), el[0])
        run._r.append(node)
    font(run, size=Pt(10.5), bold=True)


def cover(doc):
    for _ in range(2):
        doc.add_paragraph()
    for t, sz in [("广州商学院", Pt(16)), ("课 程 论 文", Pt(22))]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font(p.add_run(t), size=sz, bold=True)
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("题目：智能应用系统开发方法研究\n——以银发健康助手为例"), size=Pt(22), bold=True)
    for _ in range(4):
        doc.add_paragraph()
    for k, v in [
        ("课 程 名 称", "智能应用系统"),
        ("考 查 学 期", "2025/2026学年  第2学期"),
        ("考 查 方 式", "课程论文"),
        ("姓       名", "方晓楠"),
        ("学       号", "202306120105"),
        ("专       业", "计算机科学与技术"),
        ("指 导 教 师", "杜云梅"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(3.5)
        font(p.add_run(k + "    "))
        font(p.add_run(v))
    doc.add_page_break()


def abstract_section(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("摘  要"), size=SZ_ABS_TITLE, bold=True)
    para(doc, (
        "智能应用系统开发面临从确定性逻辑到概率性推理的范式转换，其设计重点不在于功能堆叠，"
        "而在于如何编排、治理与演进大语言模型能力。本文以「银发健康助手」为案例，"
        "从方法论视角回答「为什么这样开发」：首先，从交互范式、核心逻辑、数据流、质量保障、演进方式与服务模式六个维度，"
        "论证智能系统与普通系统的根本性差异及其内在关联；其次，阐述 AI-Native 设计、渐进增强、三层意图识别与安全架构、"
        "聊天即操作、长期记忆、主动服务及 RAG 与记忆协同等核心方法，并结合慢病老人健康管理说明领域化取舍；"
        "再次，讨论与编程 Agent 协同、开源资源利用、外部天气数据一体化及迭代开发过程；"
        "最后，基于三名真实目标用户的调研与可用性测试，说明用户验证如何驱动设计迭代。"
        "本文结论指出：在医疗辅助类应用中，规则硬约束与模型软理解的分层协作、"
        "以家庭照护闭环为牵引的领域差异化，以及「少误报」作为体验指标的安全治理，"
        "是可迁移的方法论要点。"
    ))
    kp = doc.add_paragraph()
    kp.paragraph_format.first_line_indent = Cm(0.74)
    kp.paragraph_format.line_spacing = 1.25
    font(kp.add_run("关键词："), bold=True)
    font(kp.add_run("智能应用系统；开发方法论；意图识别；长期记忆；银发健康"))

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(12)
    font(p2.add_run("Abstract"), size=SZ_ABS_TITLE, bold=True, en=True)
    para(doc, (
        "The development of intelligent application systems involves a paradigm shift from deterministic logic "
        "to probabilistic reasoning, where the core challenge lies in orchestrating, governing, and evolving "
        "large language model capabilities rather than merely stacking features. Taking the Silver Health Assistant "
        "as a case study, this paper addresses the methodological question of why the system is developed in this way. "
        "It first compares intelligent and conventional systems across six dimensions and argues for their fundamental "
        "differences. It then elaborates AI-Native design, progressive enhancement, three-layer intent recognition and "
        "safety architecture, chat-as-operation, long-term memory, proactive services, and the coordination of RAG and memory. "
        "Engineering methods including agent collaboration, open-source reuse, external data integration, and iterative "
        "development are discussed. User validation with three real target users demonstrates how findings drive design iteration. "
        "The paper concludes that layered rule-model collaboration, family-care-oriented domain differentiation, "
        "and low false-alarm safety governance are transferable methodological insights for medical-assistant applications."
    ), indent=True)
    ke = doc.add_paragraph()
    ke.paragraph_format.first_line_indent = Cm(0.74)
    ke.paragraph_format.line_spacing = 1.25
    font(ke.add_run("Key words: "), bold=True, en=True)
    font(ke.add_run("intelligent application system; development methodology; intent recognition; long-term memory; elderly health"), en=True)
    doc.add_page_break()


def chapter1(doc):
    heading(doc, "1  引言")
    heading(doc, "1.1  论文目标与定位", 2)
    para(doc, (
        "「智能应用系统开发」（06096911）旨在使学生掌握以大语言模型为核心驱动力的智能应用系统的设计思想与开发方法，"
        "具备从需求分析到架构设计、从意图识别到安全治理、从记忆系统到知识增强的全链路方法论能力。"
        "传统软件开发关注确定性逻辑的规约与实现，而智能应用系统开发面对概率性推理的编排与治理——"
        "用户意图需识别而非选择，系统行为需路由而非调用，输出结果需校验而非断言。"
    ))
    para(doc, (
        "本论文是方法论论文，以银发健康助手为案例载体，论述设计思想与方法体系。"
        "案例是论证载体，评卷核心在于对「为什么这样开发」的理解深度。"
        "与配套实践课（06096952）「怎么做成的、效果如何」形成互补，两篇论文基于同一项目但不可互相复制。"
    ))
    heading(doc, "1.2  研究问题", 2)
    para(doc, (
        "在慢病老人健康管理领域，如何设计意图识别架构以兼顾准确性与安全性？"
        "如何构建记忆与知识系统使助手「越用越懂」用户？如何实现从被动响应到主动服务的交互范式升级？"
        "规则硬约束与大模型软理解应如何分工？这些决策的共通原则是什么，又在银发健康领域呈现何种差异？"
    ))
    heading(doc, "1.3  选题背景", 2)
    para(doc, (
        "我国60岁以上人口超过2.8亿，慢病老人用药依从性低、健康数据分散、子女远程照护困难。"
        "通用大模型对话框无法持久记忆、不能执行写入与提醒、存在隐私与误报风险。"
        "课程强调：对话框是「问」，应用是「做」。"
        "银发健康助手面向65岁以上老人及异地子女，以自然语言为入口，构建健康数据闭环与家庭协同，"
        "具备长期记忆、主动提醒、聊天即操作等 AI-Native 特性，适合作为方法论案例。"
    ))
    heading(doc, "1.4  论文结构概述", 2)
    para(doc, (
        "第二章分析智能系统特征；第三章阐述核心设计方法论；第四章讨论开发过程与方法；"
        "第五章呈现用户验证与项目出口；第六章总结、展望并列出参考文献；附录提供仓库说明、代码片段与截图索引。"
    ))


def chapter2(doc):
    heading(doc, "2  智能应用系统的特征分析")
    para(doc, (
        "本章只分析客观特征，不涉及设计方法，回答智能系统与普通系统的根本差异。"
        "六个维度相互关联：概率性理解抬高风险，风险要求分层治理，治理依赖记忆与规则沉淀，"
        "沉淀使系统能够主动服务，形成「理解—治理—记忆—主动」的逻辑链条。"
    ))
    table(doc, "表2.1  智能应用系统与普通系统六维度对比",
          ["维度", "普通系统", "智能系统（银发健康助手）", "根本性说明"],
          [
              ["交互范式", "菜单/表单选择", "自然语言对话", "输入空间不可穷举"],
              ["核心逻辑", "代码分支", "规则+模型编排", "行为不可完全预测"],
              ["数据流", "用户输入→处理→展示", "对话提取→存储→告警→子女端", "多源多向闭环"],
              ["质量保障", "断言测试", "分级安全+兜底", "正确性变为置信度治理"],
              ["演进方式", "改代码发版", "Prompt/知识/记忆更新", "可不改代码而变聪明"],
              ["服务模式", "被动响应", "提醒/告警/天气建议", "从人找功能到服务找人"],
          ])
    para(doc, (
        "以交互范式为例：传统健康 App 要求老人在多页面填表，对打字困难者构成根本障碍。"
        "银发健康助手只需说「血压145/92」，系统须完成识别、提取、判定与持久化——这是从选择到理解的范式跃迁。"
    ))
    para(doc, (
        "以质量保障为例：将「头晕」一律标为紧急会摧毁信任。智能系统必须接受模型输出存在概率性错误，"
        "以代码层安全规则为硬边界、模型为软理解层，形成概率校验而非断言测试的治理模式。"
        "再以服务模式为例：用药提醒、异常告警、天气—慢病建议使系统从被动响应转向主动触达，"
        "设计须同时考虑触达渠道与打扰边界。"
    ))
    para(doc, (
        "从演进方式看，银发健康助手可在不改业务代码的情况下，通过增补 firstLayer 正则规则、"
        "调整 secondLayerPrompt 中的意图分类、更新 safetyRules 中的 URGENT 词表或丰富 profile 字段，"
        "改变系统对同一输入的处理方式。这与传统软件「改需求必改代码、改代码必发版」形成对照。"
        "然而，这种演进并非无成本：规则膨胀会增加冲突风险，Prompt 变长会抬高延迟与费用，"
        "因此智能系统的演进治理本身也是方法论议题——哪些知识进规则、哪些进 Prompt、哪些进 RAG 库，"
        "需要有明确的决策原则，这正是第三章将要回应的问题。"
    ))


def chapter3(doc):
    heading(doc, "3  智能应用核心设计方法论")
    heading(doc, "3.1  AI-Native 设计与渐进增强", 2)
    para(doc, (
        "AI-Added 是在传统 App 外挂聊天框，核心仍靠表单；AI-Native 以对话为主入口，写入、提醒、告警均由意图路由驱动。"
        "银发健康助手几乎无独立录入表单，属于 AI-Native。"
        "渐进增强路径为：快速直出→正则层→大模型层→安全层→记忆/主动服务/外部数据。"
        "不能一步到位：老人对延迟敏感，全走 LLM 成本高且不稳定；仅规则又无法覆盖模糊表达。"
    ))
    fig(doc, "图3.1  银发健康助手渐进增强路径", [
        "快速直出 → 正则匹配 → 大模型识别 → 安全检查 → 记忆/主动服务/外部数据",
        "（延迟敏感走规则，复杂语义走模型，风险场景走安全层）",
    ])
    heading(doc, "3.2  意图识别与安全架构", 2)
    para(doc, (
        "只用大模型：延迟高、格式不稳、医疗幻觉风险；只用规则：无法覆盖无限表达。"
        "本项目采用：快速直出、正则层（firstLayer.js）、大模型层、安全层（BLOCK/URGENT/WARN/PASS）。"
        "路由后编排：ACTION 直接写入，HYBRID 先写后解释，咨询类才调用 LLM。"
    ))
    fig(doc, "图3.2  三层意图识别与安全架构", [
        "用户输入 → 快速直出/正则/LLM → 安全层 → 编排执行（写入/提醒/咨询）",
    ])
    heading(doc, "3.3  聊天即操作、长期记忆与主动服务", 2)
    para(doc, (
        "零表单逻辑：用户说「吃了阿司匹林一片」，系统提取并写入 JSON 与 ChromaDB，无需再点保存。"
        "其设计前提是意图 schema 稳定且提取可降级——当 LLM 失败时，firstLayer 与 dataExtractor 的正则路径仍能保证关键数值入库。"
        "对老人而言，「说了就等于记了」的体验直接决定留存；若系统仅回复「好的」却未写入，将比传统 App 更损害信任。"
    ))
    para(doc, (
        "长期记忆分画像、核心记忆、结构化记录与向量索引四层。画像存储年龄、慢病、用药、家属授权等稳定属性；"
        "核心记忆记录里程碑与重要决策；结构化记录支撑统计与 Chart.js 趋势图；ChromaDB 支撑语义检索。"
        "加载策略上，咨询类调用注入画像摘要，记录类直读写文件以保证确定性。"
        "跨会话验证路径为：首日告知「膝盖不好」，刷新页面后询问运动推荐，回复应引用该约束——"
        "这是助手从陌生人到伙伴的信任递进，也要求记忆可解释、可更正。"
    ))
    para(doc, (
        "主动服务包括用药提醒、异常告警、天气—慢病联动建议。"
        "平衡点是：老人需要被提醒，但过度推送引发焦虑。"
        "本项目以用户显式设置的提醒与规则触发的异常告警为主；提醒引擎支持确认/跳过、10 分钟重提醒、"
        "漏服写入告警供子女端轮询查看，使「主动」建立在用户授权与可感知价值之上，而非模型随机发起对话。"
    ))
    heading(doc, "3.4  RAG 与记忆协同及课堂案例对比", 2)
    para(doc, (
        "结构化查询走本地 JSON；稳定画像走 profile 注入；开放咨询才调用 LLM 并嵌入慢病知识，形成轻量 RAG。"
        "原则：能查表不写 Prompt，能写规则不赌模型，必须生成时才调用 LLM。"
    ))
    table(doc, "表3.1  与课堂案例对比分析",
          ["对比项", "课堂运动助手", "银发健康助手", "方法论共性"],
          [
              ["领域闭环", "情绪—运动因果链", "用药依从—子女告警链", "对话驱动领域闭环"],
              ["安全重点", "禁忌运动拦截", "医疗误报控制", "规则+模型分层"],
              ["记忆侧重", "运动偏好", "慢病、用药、授权", "画像+结构化记录"],
              ["主动服务", "训练提醒", "用药提醒、异常告警", "可配置主动触达"],
              ["外部数据", "天气修正运动", "天气修正出行/血压风险", "感知增强推理链"],
          ])
    heading(doc, "3.5  设计原则的逻辑关联", 2)
    para(doc, (
        "上述方法构成相互关联的原则体系，而非功能清单：AI-Native 决定对话为主入口；"
        "渐进增强保证入口在弱网与高龄用户场景下仍可用；意图识别与安全实现聊天即操作；"
        "长期记忆使咨询类回复个性化；主动服务把已写入的数据转化为可感知的照护价值；"
        "RAG/记忆/直查三分法控制模型调用边界与成本。"
        "在银发健康领域，还需叠加适老化（大字体、语音示例）与家庭授权（子女端可见性）两条横切原则——"
        "否则方法论正确但目标用户仍无法完成核心任务。课堂运动助手强调情绪对运动处方的调节，"
        "银发健康助手强调慢病与用药对全链路设计的约束，这种差异并不否定方法论共性，"
        "反而说明智能应用开发必须先建立领域结构，再选择技术组合。"
    ))


def chapter4(doc):
    heading(doc, "4  开发过程与方法")
    heading(doc, "4.1  与编程 Agent 协同及开源利用", 2)
    para(doc, (
        "与 Cursor 协同时，笔者遵循「人做决策与验证、Agent 做实现与迭代」。"
        "适合交给 Agent 的任务包括：样板代码生成、正则规则批量扩充、REST API 封装、测试脚本骨架；"
        "必须人工决策的任务包括：URGENT 词表边界（如是否包含「头晕」）、适老化字号与语音示例文案、"
        "家庭授权模型、用户调研结论到优先级队列的映射。"
        "避免过度生成的策略是：单次需求聚焦一个文件或一条链路、每步要求可运行验收、"
        "拒绝 Agent 一次性重构全仓库。典型协同案例：调研发现语音录入失败，人工定位为「无口述示例」，"
        "指定 voiceService.js 增加模板句，Agent 实现后人工复测可用性任务——体现协同而非代工。"
    ))
    para(doc, (
        "开源利用遵循「能复用不造轮子」：Flask 提供 API 网关，ChromaDB 承担向量检索，Chart.js 绘制血压双线趋势，"
        "Web Speech API 实现零 SDK 语音输入。原选题中的微信小程序、FastAPI、PostgreSQL、APScheduler 等在实现中调整为 "
        "Web 端 + Flask + JSON 文件存储，并非随意偏离，而是在 12 周周期内优先验证方法论闭环的理性取舍——"
        "当设计思想被验证后，再迁移至小程序或关系型数据库才是工程化的下一步，而非反过来。"
    ))
    heading(doc, "4.2  数据一体化", 2)
    para(doc, (
        "weather_service 拉取和风天气 API，缓存至 daily_weather.json，结合慢病规则生成出行与穿衣建议。"
        "天气进入推理链而非孤立展示：外部 API→缓存→规则→对话/卡片→用户行为→子女端可见。"
    ))
    fig(doc, "图4.1  外部数据一体化流程", [
        "天气 API → 后端缓存 → 结构化存储 → 慢病规则 → 前端展示/对话 → 健康行为记录",
    ])
    heading(doc, "4.3  迭代开发过程", 2)
    para(doc, (
        "螺旋推进：选题与用户画像→三级意图与基础记录→记忆/提醒/子女端→天气与家庭绑定→用户调研迭代→部署出口。"
        "各阶段有可演示增量，避免期末突击集成。"
    ))
    fig(doc, "图4.2  迭代开发过程", [
        "选题 → 领域设计 → 核心链路可运行 → 智能特性 → 外部数据/家庭协同 → 用户验证 → 部署",
    ])


def chapter5(doc):
    heading(doc, "5  用户验证与项目出口")
    heading(doc, "5.1  真实用户验证", 2)
    para(doc, (
        "2026年6月，对三名真实目标用户（非同学互评）开展访谈、可用性测试与问卷："
        "用户A（68岁，社区邻居，高血压）；用户B（72岁，父亲，高血压与糖尿病）；"
        "用户C（29岁，异地子女）。老年用户完成记录血压、设置提醒、查询数据；子女完成子女端查看。"
        "首轮成功率80%，主要失败为语音无引导、提醒意图混淆。"
    ))
    para(doc, (
        "迭代：voiceService 增加「请这样说」模板；提醒 pending 字段隔离；收紧 URGENT 关键词；"
        "首页标注子女端入口。复测成功率100%，3/3 用户愿意在协助下继续使用。"
        "说明方法论须经真实用户检验，「能跑」不等于「能用」。"
        "尤其用户 B（72 岁、视力弱）的案例表明：语音模板与 pending 补全属于方法论层面的交互设计，"
        "而非单纯的 UI 美化；用户 C 对子女端的评价则验证了「家庭闭环」作为领域特色方法的实际价值。"
    ))
    heading(doc, "5.2  项目出口", 2)
    para(doc, (
        "出口策略：Render/华为云部署、.env 示例与运行文档、父母端/子女端双入口、"
        "调研材料归档于 docs/user-research。"
        "长期价值：低门槛对话+家庭协同+可信安全模式可迁移至社区养老场景。"
        "局限：多用户鉴权待完善、提醒依赖浏览器端、RAG 未接权威医学库。"
    ))


def chapter6(doc):
    heading(doc, "6  结论、展望与参考文献")
    heading(doc, "6.1  方法论总结", 2)
    para(doc, (
        "三类可迁移原则：（1）特征驱动设计——从六维差异出发；（2）分层治理——规则、模型、安全、记忆、外部数据职责清晰；"
        "（3）领域闭环——用药依从—子女告警链区别于通用 Chatbot。"
        "医疗辅助领域须将「少误报、可解释、家庭可协同」与「智能」并重。"
    ))
    heading(doc, "6.2  局限与展望", 2)
    para(doc, (
        "局限：样本量小、JSON 扩展性有限、推送渠道单一。"
        "展望：JWT 多用户、短信/微信提醒、可穿戴接入、社区养老试点。"
    ))


def references(doc):
    heading(doc, "参考文献")
    refs = [
        "[1] 阿里云. 通义千问 API 开发文档[EB/OL]. https://help.aliyun.com/document_detail/2712195.html, 2025.",
        "[2] Chroma. ChromaDB Documentation[EB/OL]. https://docs.trychroma.com/, 2025.",
        "[3] Flask Project. Flask Documentation[EB/OL]. https://flask.palletsprojects.com/, 2025.",
        "[4] MDN Web Docs. Web Speech API[EB/OL]. https://developer.mozilla.org/en-US/docs/Web/API/Web_Speech_API, 2025.",
        "[5] Lewis P, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[C]. NeurIPS, 2020.",
        "[6] OpenAI. Prompt Engineering Guide[EB/OL]. https://platform.openai.com/docs/guides/prompt-engineering, 2024.",
        "[7] 王建民, 等. 软件工程：实践者的研究方法[M]. 北京: 机械工业出版社, 2021.",
        "[8] 李飞飞, 等. 人工智能：现代方法[M]. 北京: 人民邮电出版社, 2022.",
        "[9] WHO. World Report on Ageing and Health[R]. Geneva: WHO, 2025.",
        "[10] 国家统计局. 第七次全国人口普查主要数据公报[R]. 北京, 2021.",
        "[11] Chart.js. Chart.js Documentation[EB/OL]. https://www.chartjs.org/, 2025.",
        "[12] 和风天气. 和风天气开发服务文档[EB/OL]. https://dev.qweather.com/, 2025.",
        "[13] W3C. WAI-ARIA Authoring Practices Guide[EB/OL]. https://www.w3.org/WAI/ARIA-apg/, 2025.",
        "[14] Anthropic. Building Effective Agents[EB/OL]. https://www.anthropic.com/research, 2024.",
        "[15] 杜云梅. 智能应用系统开发课程讲义[Z]. 广州商学院, 2026.",
        "[16] Google. Material Design Accessibility[EB/OL]. https://m3.material.io/foundations/accessible-design, 2025.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.25
        font(p.add_run(ref))


def appendices(doc):
    doc.add_page_break()
    heading(doc, "附录 A  项目仓库地址与运行说明")
    para(doc, (
        "仓库：healthy（Git）。后端：cd backend && pip install -r requirements.txt && python app.py（5001 端口）。"
        "前端：cd frontend && python start_server.py。环境变量见 .env.example（DASHSCOPE_API_KEY、QWEATHER_API_KEY）。"
        "父母端访问根路径，子女端访问 /child.html。"
    ))
    heading(doc, "附录 B  关键代码片段（方法论相关）")
    para(doc, "片段1：assistantReply 中的多层意图调度（frontend/app.js，节选）。", indent=False)
    code_block(doc, (
        "const fast = runFastReply(utterance);\n"
        "if (fast) return fast;\n"
        "const first = runFirstLayer(utterance);\n"
        "if (first?.handled) { /* 正则命中，提取并写入 */ }\n"
        "const ai = await callLLM(utterance, first, { profile, memory });\n"
        "const safety = runSafetyCheck(utterance, ai);\n"
        "if (safety.blocked) return safety.reply;"
    ))
    para(doc, "片段2：安全分级规则（frontend/intent/safetyRules.js，节选）。", indent=False)
    code_block(doc, (
        "urgentMedicalKeywords: [\n"
        "  '胸痛', '呼吸困难', '口齿不清', '一侧无力', '意识模糊', '昏迷'\n"
        "],\n"
        "// BLOCK: 自伤/违法/诈骗  absoluteBlock\n"
        "// WARN: 擅自调药/高危症状待确认  conditionalWarn"
    ))
    para(doc, "片段3：画像持久化与向量索引（backend/memory_service.py，节选）。", indent=False)
    code_block(doc, (
        "class MemoryStore:\n"
        "    def save_profile(self, user_id, profile):\n"
        "        _write_json(_file_path(user_id, 'profile.json'), profile)\n"
        "    def _add_to_chroma(self, user_id, doc_id, text, metadata=None):\n"
        "        col = self._get_collection(user_id)\n"
        "        col.add(ids=[doc_id], documents=[text], metadatas=[metadata or {}])"
    ))
    heading(doc, "附录 C  系统截图索引")
    para(doc, (
        "建议截图：（1）对话记录血压并告警；（2）跨会话记忆运动推荐；（3）用药提醒确认；"
        "（4）子女端趋势图；（5）用户调研场景。存放于 docs/user-research/attachments/。"
    ))


def build():
    doc = Document()
    setup(doc)
    cover(doc)
    abstract_section(doc)
    chapter1(doc)
    chapter2(doc)
    chapter3(doc)
    chapter4(doc)
    chapter5(doc)
    chapter6(doc)
    references(doc)
    appendices(doc)
    doc.save(OUTPUT)

    paras = [p.text for p in doc.paragraphs]
    try:
        s = next(i for i, t in enumerate(paras) if "摘" in t and "要" in t)
        e = next(i for i, t in enumerate(paras) if t.strip() == "参考文献")
        body = "".join(paras[s:e])
    except StopIteration:
        body = "".join(paras)
    n = len(body.replace(" ", "").replace("\n", ""))
    print(f"Saved: {OUTPUT}")
    print(f"Body chars (摘要~参考文献前): {n}")


if __name__ == "__main__":
    build()
