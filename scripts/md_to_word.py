# -*- coding: utf-8 -*-
"""Convert 理论论文.md to Word document."""

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

INPUT = Path(r"d:\21425\大三下\智能应用系统开发\healthy\理论论文.md")
OUTPUT = Path(r"d:\21425\大三下\智能应用系统开发\healthy\理论论文.docx")

FONT = "宋体"
FONT_EN = "Times New Roman"


def set_font(run, size=12, bold=False, italic=False):
    run.font.name = FONT_EN
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:eastAsia"), FONT)


def add_rich_text(paragraph, text, size=12):
    """Parse **bold** segments."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, size=size, bold=True)
        else:
            run = paragraph.add_run(part)
            set_font(run, size=size)


def setup_doc(doc):
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    sec.gutter = Cm(0.5)


def add_body_para(doc, text, indent=True, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_rich_text(p, text, size=size)
    return p


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    sizes = {1: 16, 2: 14, 3: 12, 4: 12}
    run = p.add_run(text.strip())
    set_font(run, size=sizes.get(level, 12), bold=True)
    return p


def parse_table_row(line):
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def is_table_sep(line):
    return bool(re.match(r"^\|[\s\-:|]+\|$", line.strip()))


def add_table(doc, rows):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(cols):
            val = row[ci] if ci < len(row) else ""
            cell = table.rows[ri].cells[ci]
            cell.text = val
            for para in cell.paragraphs:
                for run in para.runs:
                    set_font(run, bold=(ri == 0))
    doc.add_paragraph()


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.left_indent = Cm(0.5)
    text = "\n".join(lines)
    run = p.add_run(text)
    set_font(run, size=10.5)
    run.font.name = "Consolas"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:eastAsia"), "Consolas")
    doc.add_paragraph()


def convert():
    text = INPUT.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    setup_doc(doc)

    i = 0
    in_code = False
    code_lines = []
    table_buf = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # code block
        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # table
        if stripped.startswith("|") and stripped.endswith("|"):
            if is_table_sep(stripped):
                i += 1
                continue
            table_buf.append(parse_table_row(stripped))
            # peek next
            if i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
                i += 1
                continue
            add_table(doc, table_buf)
            table_buf = []
            i += 1
            continue

        if table_buf:
            add_table(doc, table_buf)
            table_buf = []

        # headings
        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            add_heading(doc, m.group(2), level)
            i += 1
            continue

        # horizontal rule / empty
        if stripped in ("", "---", "***"):
            i += 1
            continue

        # ordered list
        m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.25
            p.paragraph_format.first_line_indent = Cm(0.74)
            add_rich_text(p, f"{m.group(1)}. {m.group(2)}")
            i += 1
            continue

        # bullet list (- or *)
        m = re.match(r"^[-*]\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.25
            p.paragraph_format.first_line_indent = Cm(0.74)
            add_rich_text(p, "• " + m.group(1))
            i += 1
            continue

        # blockquote-like bold line starting with **
        if stripped.startswith("**") and stripped.endswith("**") and "：" not in stripped and "：" not in stripped:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.25
            p.paragraph_format.first_line_indent = Cm(0.74)
            add_rich_text(p, stripped)
            i += 1
            continue

        # normal paragraph; merge consecutive non-empty lines
        para_lines = [stripped]
        while i + 1 < len(lines):
            nxt = lines[i + 1].strip()
            if (
                not nxt
                or nxt.startswith("#")
                or nxt.startswith("|")
                or nxt.startswith("```")
                or nxt == "---"
                or re.match(r"^(\d+\.|[-*])\s", nxt)
            ):
                break
            para_lines.append(nxt)
            i += 1
        add_body_para(doc, " ".join(para_lines))
        i += 1

    if table_buf:
        add_table(doc, table_buf)

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    convert()
